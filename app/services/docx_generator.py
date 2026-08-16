"""
مولّد ملف Word (.docx) قابل للتعديل باستخدام python-docx
يدعم لغتين: العربية (اتجاه RTL) والإنجليزية (اتجاه LTR) بحسب cv_data["_language"]
يُستخدم بعد موافقة الأدمن على الدفعة
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_LABELS = {
    "ar": {
        "summary": "نبذة مهنية",
        "experience": "الخبرات العملية",
        "education": "التعليم",
        "skills": "المهارات",
        "languages": "اللغات",
    },
    "en": {
        "summary": "Professional Summary",
        "experience": "Experience",
        "education": "Education",
        "skills": "Skills",
        "languages": "Languages",
    },
}


def _set_paragraph_direction(paragraph, is_arabic: bool) -> None:
    """يضبط اتجاه الفقرة: RTL للعربية أو LTR (الافتراضي) للإنجليزية."""
    if is_arabic:
        pPr = paragraph._p.get_or_add_pPr()
        bidi = pPr.makeelement(qn("w:bidi"), {})
        pPr.append(bidi)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_heading(doc: Document, text: str, is_arabic: bool, size: int = 14, color: RGBColor | None = None):
    p = doc.add_paragraph()
    _set_paragraph_direction(p, is_arabic)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if is_arabic:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.makeelement(qn("w:rFonts"), {qn("w:cs"): "Traditional Arabic"})
        rPr.append(rFonts)
    if color:
        run.font.color.rgb = color
    return p


def _add_paragraph(doc: Document, text: str, is_arabic: bool, size: int = 11, bold: bool = False):
    p = doc.add_paragraph()
    _set_paragraph_direction(p, is_arabic)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if is_arabic:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.makeelement(qn("w:rFonts"), {qn("w:cs"): "Traditional Arabic"})
        rPr.append(rFonts)
    return p


def generate_cv_docx(cv_data: dict[str, Any]) -> bytes:
    """يولّد ملف DOCX بالكامل في الذاكرة."""
    language = cv_data.get("_language", "ar")
    is_arabic = language == "ar"
    labels = _LABELS.get(language, _LABELS["ar"])

    doc = Document()

    # اتجاه المستند بالكامل من اليمين لليسار فقط في حال كانت السيرة عربية
    if is_arabic:
        section = doc.sections[0]
        sectPr = section._sectPr
        bidi = sectPr.makeelement(qn("w:bidi"), {})
        sectPr.append(bidi)

    accent = RGBColor(0x12, 0x3A, 0x5E)

    _add_heading(doc, cv_data.get("full_name", ""), is_arabic, size=22, color=accent)
    if cv_data.get("title"):
        _add_paragraph(doc, cv_data["title"], is_arabic, size=13, bold=True)

    contact = cv_data.get("contact") or {}
    contact_line = "  |  ".join(
        filter(None, [contact.get("phone"), contact.get("email"), contact.get("location")])
    )
    if contact_line:
        _add_paragraph(doc, contact_line, is_arabic, size=10)

    doc.add_paragraph()

    if cv_data.get("summary"):
        _add_heading(doc, labels["summary"], is_arabic, size=14, color=accent)
        _add_paragraph(doc, cv_data["summary"], is_arabic)

    experience = cv_data.get("experience") or []
    if experience:
        _add_heading(doc, labels["experience"], is_arabic, size=14, color=accent)
        for exp in experience:
            header = f"{exp.get('role', '')} — {exp.get('company', '')} ({exp.get('period', '')})"
            _add_paragraph(doc, header, is_arabic, size=12, bold=True)
            for bullet in exp.get("bullets", []):
                _add_paragraph(doc, f"•  {bullet}", is_arabic, size=10.5)

    education = cv_data.get("education") or []
    if education:
        _add_heading(doc, labels["education"], is_arabic, size=14, color=accent)
        for edu in education:
            line = f"{edu.get('degree', '')} — {edu.get('institution', '')} ({edu.get('period', '')})"
            _add_paragraph(doc, line, is_arabic)

    skills = cv_data.get("skills") or []
    if skills:
        _add_heading(doc, labels["skills"], is_arabic, size=14, color=accent)
        _add_paragraph(doc, "  •  ".join(skills), is_arabic)

    languages = cv_data.get("languages") or []
    if languages:
        _add_heading(doc, labels["languages"], is_arabic, size=14, color=accent)
        _add_paragraph(doc, "  •  ".join(languages), is_arabic)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
