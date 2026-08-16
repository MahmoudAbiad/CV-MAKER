import os
import io
import json
import asyncio
from typing import Optional, List
from pydantic import BaseModel, Field

from aiogram import Bot, Dispatcher, F, types
from aiogram.filters import CommandStart
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.redis import RedisStorage
from aiogram.types import (
    InlineKeyboardMarkup,
    InlineKeyboardButton,
    FSInputFile,
    CallbackQuery,
    Message
)
from aiogram.webhook.aiohttp_server import SimpleRequestHandler, setup_application

import redis.asyncio as aioredis
from aiohttp import web
import google.generativeai as genai
from PIL import Image

import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# --- المتغيرات البيئية والإعدادات ---
BOT_TOKEN = os.getenv("BOT_TOKEN", "YOUR_BOT_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "YOUR_GEMINI_KEY")
ADMIN_GROUP_ID = int(os.getenv("ADMIN_GROUP_ID", "0"))
REDIS_URL = os.getenv("REDIS_URL", "rediss://default:password@host:port")
RENDER_EXTERNAL_URL = os.getenv("RENDER_EXTERNAL_URL", "https://your-service-name.onrender.com")
WEBHOOK_SECRET = os.getenv("WEBHOOK_SECRET", "super_secret_token_123")
PORT = int(os.getenv("PORT", 8080))
WEBHOOK_PATH = "/webhook"

SHAM_CASH_QR_PATH = "sham_cash_qr.png"

# --- تهيئة Redis و Bot و Dispatcher ---
redis_client = aioredis.from_url(REDIS_URL, decode_responses=True)
storage = RedisStorage(redis=redis_client)

genai.configure(api_key=GEMINI_API_KEY)
bot = Bot(token=BOT_TOKEN)
dp = Dispatcher(storage=storage)

# --- دوال التعامل مع Upstash Redis ---
async def save_user_cv(user_id: int, cv_dict: dict):
    await redis_client.set(f"cv:{user_id}", json.dumps(cv_dict, ensure_ascii=False))

async def get_user_cv(user_id: int) -> Optional[dict]:
    data = await redis_client.get(f"cv:{user_id}")
    return json.loads(data) if data else None

async def create_order(user_id: int, file_id: str, ocr_summary: str) -> int:
    order_id = await redis_client.incr("orders:counter")
    order_data = {
        "user_id": str(user_id),
        "file_id": file_id,
        "ocr_summary": ocr_summary,
        "status": "PENDING"
    }
    await redis_client.hset(f"order:{order_id}", mapping=order_data)
    return order_id

async def update_order_status(order_id: int, status: str, reject_reason: str = ""):
    await redis_client.hset(f"order:{order_id}", mapping={
        "status": status,
        "reject_reason": reject_reason
    })

# --- حالات FSM ---
class CVFlow(StatesGroup):
    waiting_for_bio = State()
    waiting_for_receipt = State()
    waiting_for_reject_reason = State()

# --- Schemas لـ Gemini ---
class Experience(BaseModel):
    title: str
    company: str
    dates: str
    description: List[str]

class Education(BaseModel):
    degree: str
    institution: str
    dates: str

class CVData(BaseModel):
    is_arabic: bool = Field(description="True if the text is mostly Arabic")
    full_name: str
    email: str
    phone: str
    location: str
    linkedin: Optional[str] = None
    summary: str
    skills: List[str]
    experience: List[Experience]
    education: List[Education]
    languages: List[str]

async def extract_cv_data(user_text: str) -> dict:
    model = genai.GenerativeModel(
        model_name="gemini-1.5-flash",
        generation_config={
            "response_mime_type": "application/json",
            "response_schema": CVData
        }
    )
    prompt = f"""
    You are an expert ATS Resume Parser. Extract all details into structured JSON.
    Format clearly, polish grammar, and enforce active ATS verbs.
    
    Bio Text:
    {user_text}
    """
    response = await asyncio.to_thread(model.generate_content, prompt)
    return json.loads(response.text)

async def analyze_receipt_image(image_bytes: bytes) -> str:
    model = genai.GenerativeModel("gemini-1.5-flash")
    img = Image.open(io.BytesIO(image_bytes))
    prompt = """
    أنت مدقق مالي لإيصالات شام كاش (Sham Cash).
    افحص الصورة وأعط تقريراً موجزاً باللغة العربية:
    1. هل هو إشعار شام كاش؟
    2. المبلغ المحول (هل يطابق 20 ليرة سورية جديدة؟)
    3. رقم وتاريخ العملية
    4. حالة العملية (ناجحة / قيد المعالجة)
    """
    response = await asyncio.to_thread(model.generate_content, [prompt, img])
    return response.text.strip()

def ar_text(text: str) -> str:
    if not text:
        return ""
    if any('\u0600' <= c <= '\u06FF' for c in text):
        reshaped = arabic_reshaper.reshape(text)
        return get_display(reshaped)
    return text

# --- توليد ملفات السيرة الذاتية ---
def generate_pdf(data: dict, output_path: str):
    doc = SimpleDocTemplate(
        output_path, pagesize=letter,
        rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36
    )
    is_ar = data.get('is_arabic', False)
    styles = getSampleStyleSheet()
    align = TA_RIGHT if is_ar else TA_LEFT

    font_name = 'Helvetica'
    font_bold = 'Helvetica-Bold'
    if os.path.exists("Cairo-Regular.ttf") and os.path.exists("Cairo-Bold.ttf"):
        pdfmetrics.registerFont(TTFont('ArabicFont', 'Cairo-Regular.ttf'))
        pdfmetrics.registerFont(TTFont('ArabicFont-Bold', 'Cairo-Bold.ttf'))
        font_name = 'ArabicFont'
        font_bold = 'ArabicFont-Bold'

    name_style = ParagraphStyle('Name', parent=styles['Normal'], fontName=font_bold, fontSize=18, leading=22, alignment=TA_CENTER)
    contact_style = ParagraphStyle('Contact', parent=styles['Normal'], fontName=font_name, fontSize=9, leading=12, alignment=TA_CENTER)
    heading_style = ParagraphStyle('Heading', parent=styles['Normal'], fontName=font_bold, fontSize=12, leading=16, alignment=align, spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontName=font_name, fontSize=10, leading=13, alignment=align)
    bullet_style = ParagraphStyle('Bullet', parent=styles['Normal'], fontName=font_name, fontSize=9.5, leading=12, alignment=align, leftIndent=12 if not is_ar else 0, rightIndent=12 if is_ar else 0)

    story = []
    story.append(Paragraph(ar_text(data.get('full_name', '')), name_style))
    contact_parts = [data.get('email', ''), data.get('phone', ''), data.get('location', ''), data.get('linkedin', '')]
    story.append(Paragraph(ar_text(" | ".join([p for p in contact_parts if p])), contact_style))
    story.append(Spacer(1, 10))

    if data.get('summary'):
        sec_title = "الملخص المهني" if is_ar else "PROFESSIONAL SUMMARY"
        story.append(Paragraph(ar_text(sec_title), heading_style))
        story.append(Paragraph(ar_text(data['summary']), body_style))
        story.append(Spacer(1, 6))

    if data.get('experience'):
        sec_title = "الخبرات المهنية" if is_ar else "WORK EXPERIENCE"
        story.append(Paragraph(ar_text(sec_title), heading_style))
        for exp in data['experience']:
            header = f"<b>{exp.get('title')}</b> - {exp.get('company')} ({exp.get('dates')})"
            story.append(Paragraph(ar_text(header), body_style))
            for desc in exp.get('description', []):
                story.append(Paragraph(ar_text(f"• {desc}"), bullet_style))
        story.append(Spacer(1, 6))

    if data.get('education'):
        sec_title = "التعليم والمؤهلات" if is_ar else "EDUCATION"
        story.append(Paragraph(ar_text(sec_title), heading_style))
        for edu in data['education']:
            line = f"<b>{edu.get('degree')}</b>, {edu.get('institution')} ({edu.get('dates')})"
            story.append(Paragraph(ar_text(line), body_style))
        story.append(Spacer(1, 6))

    if data.get('skills'):
        sec_title = "المهارات" if is_ar else "SKILLS"
        story.append(Paragraph(ar_text(sec_title), heading_style))
        story.append(Paragraph(ar_text(", ".join(data['skills'])), body_style))

    doc.build(story)

def generate_docx(data: dict, output_path: str):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.5)
        s.right_margin = Inches(0.5)

    is_ar = data.get('is_arabic', False)

    p_name = doc.add_paragraph()
    p_name.alignment = 1
    r_name = p_name.add_run(data.get('full_name', ''))
    r_name.bold = True
    r_name.font.size = Pt(18)

    p_contact = doc.add_paragraph()
    p_contact.alignment = 1
    contact_parts = [data.get('email', ''), data.get('phone', ''), data.get('location', ''), data.get('linkedin', '')]
    p_contact.add_run(" | ".join([p for p in contact_parts if p])).font.size = Pt(9.5)

    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11)

    if data.get('summary'):
        add_section_header("الملخص المهني" if is_ar else "PROFESSIONAL SUMMARY")
        doc.add_paragraph(data['summary'])

    if data.get('experience'):
        add_section_header("الخبرات العملية" if is_ar else "WORK EXPERIENCE")
        for exp in data['experience']:
            p_exp = doc.add_paragraph()
            p_exp.paragraph_format.space_after = Pt(1)
            p_exp.add_run(f"{exp.get('title')} - {exp.get('company')}").bold = True
            p_exp.add_run(f" ({exp.get('dates')})").italic = True
            for desc in exp.get('description', []):
                doc.add_paragraph(desc, style='List Bullet')

    if data.get('education'):
        add_section_header("التعليم" if is_ar else "EDUCATION")
        for edu in data['education']:
            p_edu = doc.add_paragraph()
            p_edu.add_run(f"{edu.get('degree')}, {edu.get('institution')} ({edu.get('dates')})")

    if data.get('skills'):
        add_section_header("المهارات" if is_ar else "SKILLS")
        doc.add_paragraph(", ".join(data['skills']))

    doc.save(output_path)

# --- معالجات تلغرام ---

@dp.message(CommandStart())
async def start_cmd(message: Message, state: FSMContext):
    welcome_text = (
        "مرحباً بك في بوت إنشاء السيرة الذاتية الذكي المتوافق مع ATS 📄✨\n\n"
        "أرسل بياناتك كاملة في **رسالة نصية واحدة**:\n"
        "• الاسم ومعلومات التواصل (الهاتف، الإيميل، المدينة، LinkedIn).\n"
        "• نبذة مهنية عنك.\n"
        "• الخبرات الوظيفية السابقة والمهام.\n"
        "• الشهادات والجامعات.\n"
        "• المهارات واللغات.\n\n"
        "🚀 سأقوم بتحليل النص وتنسيقه فوراً لملف PDF مجاني ومتوافق مع ATS."
    )
    await message.answer(welcome_text)
    await state.set_state(CVFlow.waiting_for_bio)

@dp.message(CVFlow.waiting_for_bio, F.text)
async def process_bio(message: Message, state: FSMContext):
    status_msg = await message.answer("⏳ جاري تحليل بياناتك وتوليد الـ CV...")
    try:
        cv_data = await extract_cv_data(message.text)
        await save_user_cv(message.from_user.id, cv_data)

        pdf_path = f"cv_{message.from_user.id}.pdf"
        generate_pdf(cv_data, pdf_path)

        kb = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="💎 طلب نسخة Word قابلة للتعديل (DOCX)", callback_data="buy_docx")]
        ])

        await message.reply_document(
            document=FSInputFile(pdf_path),
            caption="✅ تم إنشاء سيرتك الذاتية بتنسيق PDF بنجاح!",
            reply_markup=kb
        )
        if os.path.exists(pdf_path):
            os.remove(pdf_path)

        await status_msg.delete()
        await state.clear()
    except Exception as e:
        await status_msg.edit_text("❌ حدث خطأ أثناء المعالجة، يرجى المحاولة لاحقاً.")

@dp.callback_query(F.data == "buy_docx")
async def buy_docx_handler(call: CallbackQuery, state: FSMContext):
    cv_data = await get_user_cv(call.from_user.id)
    if not cv_data:
        await call.answer("يرجى إنشاء السيرة الذاتية أولاً.", show_alert=True)
        return

    payment_info = (
        "💳 **شراء النسخة القابلة للتعديل (DOCX)**\n\n"
        "• السعر: **20 ليرة سورية جديدة**\n"
        "• طريقة الدفع: عبر تطبيق **شام كاش (Sham Cash)**\n\n"
        "يرجى مسح الـ QR أو التحويل لحسابنا ثم **إرسال صورة إشعار الدفع** هنا مباشرة."
    )

    if os.path.exists(SHAM_CASH_QR_PATH):
        await call.message.answer_photo(photo=FSInputFile(SHAM_CASH_QR_PATH), caption=payment_info)
    else:
        await call.message.answer(payment_info)

    await state.set_state(CVFlow.waiting_for_receipt)
    await call.answer()

@dp.message(CVFlow.waiting_for_receipt, F.photo)
async def handle_receipt(message: Message, state: FSMContext):
    user = message.from_user
    status_msg = await message.answer("🔍 جاري فحص إشعار الدفع بالذكاء الاصطناعي...")

    photo = message.photo[-1]
    file_info = await bot.get_file(photo.file_id)
    photo_bytes = await bot.download_file(file_info.file_path)
    image_data = photo_bytes.read()

    ocr_result = await analyze_receipt_image(image_data)
    order_id = await create_order(user.id, photo.file_id, ocr_result)

    admin_kb = InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="✅ موافقة وإرسال DOCX", callback_data=f"order_app_{order_id}_{user.id}"),
            InlineKeyboardButton(text="❌ رفض", callback_data=f"order_rej_{order_id}_{user.id}")
        ]
    ])

    admin_caption = (
        f"🔔 **طلب شراء ملف DOCX جديد (رقم: #{order_id})**\n\n"
        f"👤 **المستخدم:** {user.full_name} (@{user.username or 'بدون معرف'})\n"
        f"🆔 **الآيدي:** `{user.id}`\n\n"
        f"🤖 **تقرير تدقيق الإشعار آلياً:**\n{ocr_result}"
    )

    await bot.send_photo(
        chat_id=ADMIN_GROUP_ID,
        photo=photo.file_id,
        caption=admin_caption,
        reply_markup=admin_kb
    )

    await status_msg.edit_text("⏳ تم استلام إشعار الدفع بنجاح، بانتظار تأكيد الإدارة.")
    await state.clear()

@dp.callback_query(F.data.startswith("order_app_"))
async def approve_payment(call: CallbackQuery):
    _, _, order_id_str, user_id_str = call.data.split("_")
    order_id = int(order_id_str)
    user_id = int(user_id_str)

    cv_data = await get_user_cv(user_id)
    if not cv_data:
        await call.answer("بيانات المستخدم غير موجودة.", show_alert=True)
        return

    docx_path = f"cv_{user_id}.docx"
    generate_docx(cv_data, docx_path)

    try:
        await bot.send_document(
            chat_id=user_id,
            document=FSInputFile(docx_path),
            caption="🎉 تم تأكيد الدفع بنجاح! تفضل نسختك القابلة للتعديل بصيغة Word (DOCX)."
        )
        await update_order_status(order_id, "APPROVED")
        await call.message.edit_caption(
            caption=call.message.caption + "\n\n✅ **تمت الموافقة وتم تسليم الملف بنجاح.**",
            reply_markup=None
        )
        await call.answer("تمت الموافقة والتسليم.")
    except Exception as e:
        await call.answer(f"فشل الإرسال للمستخدم: {e}", show_alert=True)
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)

@dp.callback_query(F.data.startswith("order_rej_"))
async def reject_payment_prompt(call: CallbackQuery, state: FSMContext):
    _, _, order_id_str, user_id_str = call.data.split("_")
    await state.update_data(target_order_id=int(order_id_str), target_user_id=int(user_id_str))
    await state.set_state(CVFlow.waiting_for_reject_reason)
    await call.message.reply("✏️ أرسل سبب الرفض في رسالة:")
    await call.answer()

@dp.message(CVFlow.waiting_for_reject_reason, F.text)
async def process_reject_reason(message: Message, state: FSMContext):
    data = await state.get_data()
    user_id = data.get("target_user_id")
    order_id = data.get("target_order_id")
    reason = message.text

    await update_order_status(order_id, "REJECTED", reason)

    try:
        await bot.send_message(
            chat_id=user_id,
            text=f"❌ تم رفض طلب الدفع رقم #{order_id}.\n\n**السبب:** {reason}"
        )
        await message.reply("✅ تم إشعار المستخدم بالرفض.")
    except Exception as e:
        await message.reply(f"تعذر مراسلة المستخدم: {e}")

    await state.clear()

# --- إدارة دورة حياة الـ Webhook و aiohttp ---

async def on_startup(bot: Bot):
    webhook_url = f"{RENDER_EXTERNAL_URL.rstrip('/')}{WEBHOOK_PATH}"
    await bot.set_webhook(
        url=webhook_url,
        secret_token=WEBHOOK_SECRET,
        drop_pending_updates=True,
        allowed_updates=["message", "callback_query"]
    )
    print(f"🚀 Webhook configured successfully: {webhook_url}")

async def on_shutdown(bot: Bot):
    await bot.delete_webhook()
    await redis_client.aclose()
    await bot.session.close()
    print("🛑 Webhook removed and sessions closed.")

def main():
    dp.startup.register(on_startup)
    dp.shutdown.register(on_shutdown)

    app = web.Application()

    # فحص الجاهزية (Health Check) لخدمة Render
    async def health_check(request):
        return web.Response(text="Bot Webhook is healthy and running!")

    app.router.add_get("/", health_check)
    app.router.add_get("/health", health_check)

    # تسجيل معالج تحديثات تلغرام
    webhook_requests_handler = SimpleRequestHandler(
        dispatcher=dp,
        bot=bot,
        secret_token=WEBHOOK_SECRET,
    )
    webhook_requests_handler.register(app, path=WEBHOOK_PATH)

    setup_application(app, dp, bot=bot)
    web.run_app(app, host="0.0.0.0", port=PORT)

if __name__ == "__main__":
    main()
